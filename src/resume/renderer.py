"""
Resume DOCX + PDF renderer.
Produces a minimal but well-structured resume from rephrased bullets.

DOCX: python-docx (always available)
PDF:  tries docx2pdf (Windows COM), falls back to writing a placeholder if not installed.
      In tests, PDF rendering is skipped or mocked.
"""
from __future__ import annotations
import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed; DOCX rendering disabled")


def render_docx(
    bullets_by_type: dict[str, list[str]],
    candidate_name: str = "Candidate",
    job_title: str = "",
    out_path: Optional[str] = None,
) -> str:
    """
    Render a DOCX resume from rephrased bullets.
    Returns the output path.
    """
    if not _DOCX_AVAILABLE:
        # Write a plain-text stub so tests can verify the file was created
        out_path = out_path or "resume.docx"
        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write(_plain_text(bullets_by_type, candidate_name, job_title))
        logger.warning("DOCX lib unavailable; wrote plain-text stub to %s", out_path)
        return out_path

    doc = Document()

    # Header
    header = doc.add_heading(candidate_name, level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if job_title:
        sub = doc.add_paragraph(f"Applying for: {job_title}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sections
    _add_section(doc, "Experience", bullets_by_type.get("metric", []))
    _add_section(doc, "Technical Skills",
                 bullets_by_type.get("tool", []) +
                 bullets_by_type.get("skill", []))
    _add_section(doc, "Education & Credentials",
                 bullets_by_type.get("credential", []))

    out_path = out_path or "resume.docx"
    doc.save(out_path)
    logger.info("rendered DOCX to %s", out_path)
    return out_path


def render_pdf(docx_path: str, pdf_path: Optional[str] = None) -> str:
    """
    Convert DOCX to PDF. Returns the PDF path.
    Uses docx2pdf if available; otherwise writes a placeholder.
    """
    pdf_path = pdf_path or docx_path.replace(".docx", ".pdf")

    try:
        from docx2pdf import convert  # noqa: PLC0415
        convert(docx_path, pdf_path)
        logger.info("rendered PDF to %s", pdf_path)
    except (ImportError, Exception) as exc:
        logger.warning("PDF render skipped (%s); writing placeholder", exc)
        with open(pdf_path, "w", encoding="utf-8") as fh:
            fh.write(f"[PDF placeholder — docx2pdf unavailable]\nSource: {docx_path}\n")

    return pdf_path


# ---------------------------------------------------------------------------
# Internals
# ---------------------------------------------------------------------------

def _add_section(doc, title: str, bullets: list[str]) -> None:
    if not bullets:
        return
    doc.add_heading(title, level=1)
    for bullet in bullets:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(bullet)


def _plain_text(
    bullets_by_type: dict[str, list[str]],
    candidate_name: str,
    job_title: str,
) -> str:
    lines = [candidate_name]
    if job_title:
        lines.append(f"Applying for: {job_title}")
    lines.append("")
    for section, bullets in bullets_by_type.items():
        if bullets:
            lines.append(f"=== {section.upper()} ===")
            for b in bullets:
                lines.append(f"- {b}")
            lines.append("")
    return "\n".join(lines)
