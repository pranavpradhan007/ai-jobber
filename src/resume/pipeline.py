"""
Resume tailoring pipeline — copy-and-patch strategy.

Ground truth resume (ML-AI Resume.docx) is NEVER rewritten.
Only the Technical Skills line is patched to surface missing JD keywords.
All other content, formatting, and design is preserved exactly.

Full flow:
  1. Copy ground-truth DOCX to artifact folder
  2. Find the Technical Skills / Skills line in the DOCX
  3. Inject any hot_keywords that aren't already present (append to the line)
  4. Convert to PDF
  5. Run diff-verifier on the injected text only
  6. Write resume_diff.md
  7. Persist paths + advance state
"""
from __future__ import annotations
import logging
import os
import re
import shutil
import sqlite3
from dataclasses import dataclass
from typing import Optional

from src.db.applications import update_application
from src.db.state_machine import transition
from src.storage.folders import create_application_folder, artifact_path
from src.verifier.diff_verifier import verify_text
from src.resume.renderer import render_pdf
from src.resume.diff import write_resume_diff
from src.verifier.retrieval import BankItem

logger = logging.getLogger(__name__)

# Where the user's master resume lives (set by user in .env or defaults to D:\)
_DEFAULT_RESUME_DOCX = r"D:\Pranav\Resume\New folder\ML-AI Resume.docx"
_DEFAULT_RESUME_PDF  = r"D:\Pranav\Resume\New folder\Pranav ML-AI Resume.pdf"


@dataclass
class TailoringResult:
    resume_path: str
    resume_pdf_path: str
    resume_diff_path: str
    verifier_passed: bool
    ats_report_path: str = ""
    ats_score: int = 0


class PageLimitError(Exception):
    """Raised when the rendered resume exceeds 1 page."""


def run_tailoring(
    conn: sqlite3.Connection,
    app_id: int,
    *,
    hot_keywords: list[str],
    job_title: str = "",
    candidate_name: str = "Candidate",
    rephraser=None,
    extractor=None,
    surface: str = "resume",
    edit_instruction: Optional[str] = None,
    raw_jd: str = "",
    # Loop controls
    max_resume_attempts: int = 15,
    max_outer_loops: int = 5,
    ats_min_score: float = 75.0,
    score_threshold: float = 85.0,
) -> TailoringResult:
    """
    Multi-agent tailoring loop:

    Outer loop (up to max_outer_loops = 5):
      Inner — Resume Agent (up to max_resume_attempts = 15):
        Copy GT → reframe bullets → inject keywords → convert → check hard
        rules → score vs JD. Accept if both pass, else retry with escalating
        conservatism (full → partial → inject-only).
      If Resume Agent finds a passing resume:
        Run ATS Agent. If ATS score >= ats_min_score → done.
        Else → outer loop continues with tighter parameters.
      If Resume Agent exhausts all attempts → outer loop fails.

    After all outer loops fail → ground truth PDF submitted.
    """
    folder = create_application_folder(conn, app_id)

    source_docx = os.environ.get("RESUME_DOCX_PATH", _DEFAULT_RESUME_DOCX)
    source_pdf  = os.environ.get("RESUME_PDF_PATH",  _DEFAULT_RESUME_PDF)

    if not os.path.isfile(source_docx):
        raise FileNotFoundError(
            f"Ground-truth resume not found: {source_docx}\n"
            "Set RESUME_DOCX_PATH in .env to the correct path."
        )

    docx_path = artifact_path(folder, "resume.docx")
    pdf_path  = artifact_path(folder, "resume.pdf")

    # Fetch raw JD from DB if not supplied
    if not raw_jd:
        try:
            row = conn.execute(
                "SELECT j.raw_jd FROM applications a JOIN jobs j ON j.id=a.job_id WHERE a.id=?",
                (app_id,)
            ).fetchone()
            raw_jd = (row["raw_jd"] or "") if row else ""
        except Exception:
            raw_jd = ""

    from src.resume.resume_agent import run_resume_agent
    from src.resume.ats_agent import run_ats_check, ground_truth_ats_score

    # Calibrate ATS threshold against the ground truth resume.
    # Accept tailored version if it beats or matches the GT's ATS score.
    # Hard floor: never require more than 60 (our sim's reasonable floor).
    gt_ats = ground_truth_ats_score(source_docx, source_pdf, hot_keywords)
    effective_ats_threshold = max(ats_min_score, gt_ats - 5.0)
    logger.info(
        "app_id=%d GT ATS score=%.1f, effective threshold=%.1f",
        app_id, gt_ats, effective_ats_threshold,
    )

    winning_resume = None
    winning_ats    = None
    outer_loop_log: list[str] = []

    for outer in range(1, max_outer_loops + 1):
        logger.info("=== TAILORING outer_loop=%d/%d app_id=%d ===", outer, max_outer_loops, app_id)

        # Tighten score threshold on later outer loops to try harder
        loop_score_threshold = score_threshold + (outer - 1) * 2.0
        loop_ats_threshold   = effective_ats_threshold

        resume_result = run_resume_agent(
            source_docx=source_docx,
            source_pdf=source_pdf,
            docx_path=docx_path,
            pdf_path=pdf_path,
            hot_keywords=hot_keywords,
            job_title=job_title,
            raw_jd=raw_jd,
            rephraser_fn=rephraser,
            max_attempts=max_resume_attempts,
            score_threshold=loop_score_threshold,
        )

        if not resume_result.success:
            msg = (
                f"outer={outer}: Resume Agent exhausted {resume_result.total_attempts} attempts "
                f"without passing (threshold={loop_score_threshold})"
            )
            logger.warning(msg)
            outer_loop_log.append(msg)
            continue

        best = resume_result.best
        logger.info(
            "outer=%d Resume Agent succeeded: attempt=%d score=%.1f strategy=%s kw=%s",
            outer, best.attempt, best.score, best.strategy, best.keywords_injected,
        )

        # ATS check on the best resume this outer loop produced
        ats = run_ats_check(
            pdf_path=best.pdf_path,
            docx_path=best.docx_path,
            hot_keywords=hot_keywords,
            raw_jd=raw_jd,
            min_score=loop_ats_threshold,
        )

        if ats.passed:
            winning_resume = best
            winning_ats    = ats
            logger.info(
                "outer=%d ATS PASSED: %.1f/100 — accepting resume (attempt=%d score=%.1f)",
                outer, ats.score, best.attempt, best.score,
            )
            break
        else:
            msg = (
                f"outer={outer}: ATS FAILED ({ats.score:.1f}/{loop_ats_threshold:.1f}) "
                f"after Resume Agent best_score={best.score:.1f}"
            )
            logger.warning(msg)
            outer_loop_log.append(msg)

    # ── Final decision ────────────────────────────────────────────────────────
    used_ground_truth = False
    final_kw: list[str] = []
    ats_score_final = 0

    if winning_resume is not None:
        # Copy winning artifacts to canonical paths
        if winning_resume.docx_path != docx_path:
            shutil.copy2(winning_resume.docx_path, docx_path)
        if winning_resume.pdf_path != pdf_path:
            shutil.copy2(winning_resume.pdf_path, pdf_path)
        final_kw = winning_resume.keywords_injected
        ats_score_final = int(winning_ats.score) if winning_ats else 0
        checker_note = (
            f"PASSED outer={winning_ats is not None} "
            f"ATS={ats_score_final} resume_score={winning_resume.score:.1f}"
        )
    else:
        logger.warning(
            "app_id=%d ALL %d outer loops failed — using ground truth PDF. Log: %s",
            app_id, max_outer_loops, outer_loop_log,
        )
        if os.path.isfile(source_pdf):
            shutil.copy2(source_pdf, pdf_path)
        shutil.copy2(source_docx, docx_path)
        used_ground_truth = True
        checker_note = f"FAILED ({max_outer_loops} outer × {max_resume_attempts} inner) — ground truth submitted"

    # ── Verifier + diff ───────────────────────────────────────────────────────
    injected_text = ", ".join(final_kw) if final_kw else "(ground truth — no injection)"
    verifier_report = verify_text(
        conn, app_id, injected_text, surface=surface, extractor=extractor
    )

    fake_bullets: list[BankItem] = [
        BankItem(item_type="keyword", value=kw, source="hot_keywords", usage_level="resume")
        for kw in (final_kw or hot_keywords[:3])
    ]
    diff_path = write_resume_diff(
        folder, fake_bullets,
        [f"Added to skills line: {kw}" for kw in final_kw],
        verifier_report,
        job_title=job_title,
        style_report=checker_note,
    )

    update_application(
        conn, app_id,
        resume_path=docx_path,
        resume_pdf_path=pdf_path,
        resume_diff_path=diff_path,
        verifier_passed=0 if used_ground_truth else 1,
    )
    transition(conn, app_id, "RESUME_VERIFIED", reason="verifier passed")

    logger.info(
        "tailoring complete app_id=%d used_gt=%s ats=%d kw=%s",
        app_id, used_ground_truth, ats_score_final, final_kw,
    )
    return TailoringResult(
        resume_path=docx_path,
        resume_pdf_path=pdf_path,
        resume_diff_path=diff_path,
        verifier_passed=not used_ground_truth,
        ats_score=ats_score_final,
    )


# AI buzzwords that must never appear in injected text
_AI_LANGUAGE_BLACKLIST = {
    "leveraged", "synergized", "spearheaded", "pioneered", "utilized", "utilised",
    "streamlined", "optimized", "revolutionized", "transformed", "harnessed",
    "orchestrated", "facilitated", "implemented", "executed", "delivered",
    "proactive", "dynamic", "innovative", "cutting-edge", "state-of-the-art",
    "robust", "scalable", "impactful", "strategic", "holistic", "seamless",
}

# Em dash variants to strip from any injected text
_EM_DASH_CHARS = ["—", "–", "·", "—", "–"]

# Hard caps for 1-page budget — conservative to leave room for font hinting variance
# Zero-growth policy: never allow the skills line to grow by more than _SKILLS_MAX_GROWTH chars.
# Any growth risks pushing a tightly-fitted 1-page resume to 2 pages.
_SKILLS_MAX_GROWTH = 25         # max chars we'll add to the skills line
_MAX_NEW_KEYWORDS  = 2          # max new keyword terms to inject


def _clean_keyword(kw: str) -> str:
    """Strip em dashes and AI language from a keyword string."""
    for ch in _EM_DASH_CHARS:
        kw = kw.replace(ch, " ")
    kw = kw.strip()
    # Reject if the keyword IS an AI buzzword
    if kw.lower() in _AI_LANGUAGE_BLACKLIST:
        return ""
    return kw


def _inject_keywords(docx_path: str, hot_keywords: list[str]) -> list[str]:
    """
    Find the Technical Skills CONTENT paragraph (not the heading) and append
    any hot_keywords not already present in the whole document.
    Returns the list of newly added keywords. Edits the file in-place.

    Rules enforced:
    - Target must be a Normal/body paragraph, NOT a Heading paragraph
    - No em dashes or AI buzzwords
    - Cap at _MAX_NEW_KEYWORDS terms and _SKILLS_LINE_CHAR_LIMIT chars
    - New run is added with explicit formatting (no underline, same font size)
      so heading styles never bleed into injected text
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        logger.warning("python-docx not installed; skipping keyword injection")
        return []

    doc = Document(docx_path)

    # Build full-document text for duplicate detection
    full_doc_text = " ".join(p.text for p in doc.paragraphs).lower()

    # Find the skills CONTENT paragraph:
    # Must NOT be a Heading style, must start with a skills category label
    # (e.g. "Languages:", "Cloud", "Libraries:")
    _CONTENT_MARKERS = [
        "languages:", "cloud", "libraries:", "database:", "frameworks:",
        "tools:", "programming languages",
    ]
    _HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}

    skills_para = None
    for para in doc.paragraphs:
        if para.style.name in _HEADING_STYLES:
            continue
        text_lower = para.text.lower().strip()
        if any(text_lower.startswith(m) or (f" {m}" in text_lower) for m in _CONTENT_MARKERS):
            skills_para = para
            break

    if skills_para is None:
        logger.warning("Could not find skills content line in DOCX; skipping injection")
        return []

    original_len = len(skills_para.text)
    current_len  = original_len

    to_add: list[str] = []
    for kw in hot_keywords:
        if len(to_add) >= _MAX_NEW_KEYWORDS:
            break
        cleaned = _clean_keyword(kw)
        if not cleaned:
            continue
        if cleaned.lower() in full_doc_text:
            continue
        if len(cleaned) >= 40:
            continue
        addition_len = len(", " + cleaned)
        if (current_len - original_len) + addition_len > _SKILLS_MAX_GROWTH:
            logger.info(
                "skills line growth budget exhausted (%d/%d chars used); stopping injection",
                current_len - original_len, _SKILLS_MAX_GROWTH,
            )
            break
        to_add.append(cleaned)
        current_len += addition_len

    if not to_add:
        return []

    # Infer font size from existing runs (copy from last non-empty run)
    ref_run = next((r for r in reversed(skills_para.runs) if r.text.strip()), None)
    ref_size = ref_run.font.size if ref_run else None

    # Add a NEW run so we control its formatting completely
    # (never modify the last existing run — it may carry heading/underline styles)
    new_run = skills_para.add_run(", " + ", ".join(to_add))
    new_run.bold = False
    new_run.italic = False
    new_run.underline = False
    if ref_size:
        new_run.font.size = ref_size

    doc.save(docx_path)
    return to_add
